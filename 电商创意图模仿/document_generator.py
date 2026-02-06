# -*- coding: utf-8 -*-
"""
Word 文档生成模块
将竞品分析结果整理成格式化的创意图工单文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import datetime
from PIL import Image
import io


def create_brief_document(
    results_list: list,
    output_folder: str,
    product_name: str = "雪叶红口服液",
    author_name: str = "阳宏"
) -> str:
    """
    生成创意图工单 Word 文档
    
    Args:
        results_list: 分析结果列表，每项包含 image_path 和 copywriting
        output_folder: 输出文件夹路径
        product_name: 产品名称
        author_name: 作者名称
    
    Returns:
        str: 生成的文档路径
    """
    # 确保输出文件夹存在
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 生成文件名：YYYYMMDD-电商-{product_name}-{author_name}.docx
    today = datetime.now().strftime("%Y%m%d")
    time_suffix = datetime.now().strftime("%H%M%S")
    base_filename = f"{today}-电商-{product_name}-{author_name}"
    doc_path = output_path / f"{base_filename}.docx"
    
    # 如果文件已存在且被占用，使用带时间戳的文件名
    if doc_path.exists():
        doc_path = output_path / f"{base_filename}-{time_suffix}.docx"
    
    # 创建文档
    doc = Document()
    
    # 设置文档样式
    _setup_styles(doc)
    
    # 添加标题
    title = doc.add_heading("创意图工单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    _add_basic_info(doc, product_name, len(results_list))
    
    # 添加分隔线
    doc.add_paragraph("─" * 50)
    
    # 循环添加每组结果
    for i, result in enumerate(results_list, 1):
        _add_result_section(doc, result, i, len(results_list))
    
    # 保存文档（处理可能的权限错误）
    try:
        doc.save(str(doc_path))
    except PermissionError:
        # 文件被占用，使用新的时间戳文件名
        doc_path = output_path / f"{base_filename}-{time_suffix}.docx"
        doc.save(str(doc_path))
    
    print(f"[完成] 文档已保存: {doc_path}")
    
    return str(doc_path)


def _setup_styles(doc: Document):
    """设置文档样式"""
    # 设置正文默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)


def _add_basic_info(doc: Document, product_name: str, count: int):
    """添加基本信息段落"""
    info_para = doc.add_paragraph()
    
    # 添加信息行
    info_lines = [
        f"图片类别：创意图",
        f"尺寸：800×800",
        f"数量：{count}张",
        f"产品：{product_name}",
        f"日期：{datetime.now().strftime('%Y年%m月%d日')}"
    ]
    
    for line in info_lines:
        run = info_para.add_run(line + "\n")
        run.font.size = Pt(12)


def _add_result_section(doc: Document, result: dict, index: int, total: int):
    """
    添加单个结果区块（精简版，供设计师使用）
    
    Args:
        doc: Word 文档对象
        result: 包含 image_path 和 copywriting 的字典
        index: 当前索引
        total: 总数
    """
    # 不添加编号标题，直接插入图片
    
    # 插入竞品参考图（左对齐，200px宽度）
    image_path = result.get("image_path", "")
    if image_path and Path(image_path).exists():
        try:
            # 图片段落 - 左对齐
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 固定宽度为 200px ≈ 2.08 英寸
            display_width = 2.08  # 200px at 96 DPI
            
            run = para.add_run()
            run.add_picture(image_path, width=Inches(display_width))
            
        except Exception as e:
            doc.add_paragraph(f"[图片加载失败: {e}]")
    
    # 直接添加文案内容
    copywriting = result.get("copywriting", "")
    if copywriting:
        # 分段处理文案
        paragraphs = copywriting.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.2
    
    # 添加简单的分隔线（只有一行，不分页）
    if index < total:
        doc.add_paragraph("─" * 40)


def create_simple_brief(
    results_list: list,
    output_folder: str,
    product_name: str = "雪叶红口服液"
) -> str:
    """
    创建简化版工单（仅文案，无图片）
    
    用于快速预览或网络发送
    """
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    today = datetime.now().strftime("%Y%m%d")
    filename = f"{today}-{product_name}-文案汇总.txt"
    txt_path = output_path / filename
    
    lines = [
        f"创意图文案汇总",
        f"产品：{product_name}",
        f"数量：{len(results_list)}张",
        f"日期：{datetime.now().strftime('%Y年%m月%d日')}",
        "=" * 50,
        ""
    ]
    
    for i, result in enumerate(results_list, 1):
        lines.append(f"【第 {i} 张】{result.get('image_name', '')}")
        lines.append("-" * 30)
        lines.append(result.get('copywriting', ''))
        lines.append("")
        lines.append("=" * 50)
        lines.append("")
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    print(f"[完成] 文案汇总已保存: {txt_path}")
    return str(txt_path)


if __name__ == "__main__":
    # 测试代码
    test_results = [
        {
            "image_path": "test.jpg",
            "image_name": "test.jpg",
            "copywriting": "主标题：宫寒调理专家\n\n副标题：7天感受温暖\n\n正文：采用八大珍贵中草药配方..."
        }
    ]
    
    try:
        doc_path = create_brief_document(
            test_results,
            "output",
            product_name="雪叶红口服液",
            author_name="阳宏"
        )
        print(f"测试文档已生成: {doc_path}")
    except Exception as e:
        print(f"测试失败: {e}")
