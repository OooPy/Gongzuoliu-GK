# -*- coding: utf-8 -*-
"""
Word 工单文档生成模块
将竞品分析结果整理成格式化的设计工单文档
严格遵循输出规范：禁止页码、模块标题、视觉风格建议
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
from datetime import datetime
from PIL import Image


def create_work_order(
    detail_results: list,
    main_results: list,
    output_folder: str,
    product_name: str = "雪叶红口服液",
    author_name: str = "阳宏"
) -> str:
    """
    生成完整的设计工单 Word 文档
    
    Args:
        detail_results: 详情页分析结果列表
        main_results: 主图分析结果列表
        output_folder: 输出文件夹
        product_name: 产品名称
        author_name: 负责人
        
    Returns:
        str: 生成的文档路径
    """
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 生成文件名：YYYYMMDD-电商-产品名-负责人.docx
    today = datetime.now().strftime("%Y%m%d")
    time_suffix = datetime.now().strftime("%H%M%S")
    base_filename = f"{today}-电商-{product_name}-{author_name}"
    doc_path = output_path / f"{base_filename}.docx"
    
    # 处理文件已存在的情况
    if doc_path.exists():
        doc_path = output_path / f"{base_filename}-{time_suffix}.docx"
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    _setup_styles(doc)
    
    # 添加标题
    title = doc.add_heading(f"{today}-电商-{product_name}-{author_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加规格信息
    _add_spec_info(doc, detail_results, main_results, product_name)
    
    # 添加分隔线
    doc.add_paragraph("─" * 60)
    
    # 添加主图部分（优先输出）
    if main_results:
        for result in main_results:
            _add_content_section(doc, result)
    
    # 添加详情页部分
    if detail_results:
        for result in detail_results:
            _add_content_section(doc, result)
    
    # 保存文档
    try:
        doc.save(str(doc_path))
    except PermissionError:
        doc_path = output_path / f"{base_filename}-{time_suffix}.docx"
        doc.save(str(doc_path))
    
    print(f"\n[完成] 工单文档已保存: {doc_path}")
    return str(doc_path)


def _setup_styles(doc: Document):
    """设置文档默认样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _add_spec_info(doc: Document, detail_results: list, main_results: list, product_name: str):
    """添加规格信息（严格按照规范格式）"""
    
    detail_count = len(detail_results)  # 详情页屏数
    main_count = len(main_results)      # 主图张数
    total_count = detail_count + main_count * 2  # 总数量
    
    info_para = doc.add_paragraph()
    
    # 规格信息行
    info_lines = [
        f"图片类别：详情页(需要PSD)+主图",
        f"尺寸：详情页宽度 1440（{detail_count} 屏），高度不限。主图 1440×1440（{main_count} 张）+ 1440×1920（{main_count} 张）。",
        f"数量：{total_count} 张",
        f"产品：{product_name}"
    ]
    
    for line in info_lines:
        run = info_para.add_run(line + "\n")
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'


def _add_content_section(doc: Document, result: dict):
    """
    添加单条内容记录（图片 + 文案）
    严格遵循规范：禁止页码、模块标题、视觉风格建议
    """
    
    # 插入竞品参考图（左对齐，200px 宽度）
    image_path = result.get("image_path", "")
    if image_path and Path(image_path).exists():
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 固定宽度 200px ≈ 2.08 英寸 (96 DPI)
            display_width = 2.08
            
            run = para.add_run()
            run.add_picture(image_path, width=Inches(display_width))
            
        except Exception as e:
            error_para = doc.add_paragraph(f"[图片加载失败: {e}]")
            error_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加文案内容（紧随图片，左对齐）
    copywriting = result.get("copywriting", "")
    if copywriting:
        # 分段处理
        paragraphs = copywriting.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.space_after = Pt(0)
    
    # 添加简洁分隔（不分页）
    doc.add_paragraph()
    doc.add_paragraph("─" * 40)


def create_ab_test_orders(
    detail_results_a: list,
    detail_results_b: list,
    main_results_a: list,
    main_results_b: list,
    output_folder: str,
    product_name: str = "雪叶红口服液",
    author_name: str = "阳宏"
) -> tuple:
    """
    生成 A/B 测试版本的工单文档
    
    Args:
        detail_results_a: 详情页 A 版本结果
        detail_results_b: 详情页 B 版本结果
        main_results_a: 主图 A 版本结果
        main_results_b: 主图 B 版本结果
        output_folder: 输出文件夹
        product_name: 产品名称
        author_name: 负责人
        
    Returns:
        tuple: (A版本文档路径, B版本文档路径)
    """
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    today = datetime.now().strftime("%Y%m%d")
    
    # A 版本：感性/生活场景
    doc_a_path = output_path / f"{today}-电商-{product_name}-{author_name}-A版(感性).docx"
    _create_single_order(detail_results_a, main_results_a, str(doc_a_path), product_name, author_name, "A版-感性")
    
    # B 版本：理性/数据对比
    doc_b_path = output_path / f"{today}-电商-{product_name}-{author_name}-B版(理性).docx"
    _create_single_order(detail_results_b, main_results_b, str(doc_b_path), product_name, author_name, "B版-理性")
    
    return str(doc_a_path), str(doc_b_path)


def _create_single_order(detail_results, main_results, doc_path, product_name, author_name, version_tag):
    """创建单个版本的工单文档"""
    doc = Document()
    _setup_styles(doc)
    
    today = datetime.now().strftime("%Y%m%d")
    title = doc.add_heading(f"{today}-电商-{product_name}-{author_name} ({version_tag})", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    _add_spec_info(doc, detail_results, main_results, product_name)
    doc.add_paragraph("─" * 60)
    
    for result in detail_results:
        _add_content_section(doc, result)
    
    for result in main_results:
        _add_content_section(doc, result)
    
    doc.save(doc_path)
    print(f"[完成] {version_tag} 文档已保存: {doc_path}")


if __name__ == "__main__":
    # 测试代码
    test_detail = [
        {
            "image_path": "",
            "copywriting": "3.0高阶螯合铁\n铁肽同修 内外双行\n提振内在精神气"
        }
    ]
    
    test_main = [
        {
            "image_path": "",
            "copywriting": "孕妈必备好气色\n8大珍贵成分 科学配比\n国食健字号 安全放心"
        }
    ]
    
    try:
        doc_path = create_work_order(
            test_detail,
            test_main,
            "output",
            product_name="雪叶红口服液",
            author_name="阳宏"
        )
        print(f"测试文档已生成: {doc_path}")
    except Exception as e:
        print(f"测试失败: {e}")
