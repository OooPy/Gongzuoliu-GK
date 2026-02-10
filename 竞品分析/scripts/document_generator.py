# -*- coding: utf-8 -*-
"""
Word 文档生成模块
将竞品分析结果输出为结构化的 Word 分析报告

报告结构（9大维度 + 人群画像 + 综合评估）：
1. 产品概况
2. 5张主图卖点与视觉策略分析
3. 详情页表达逻辑分析
4. 规格设计分析
5. 活动与促销策略
6. 用户评价分析
7. 近30日销售数据分析
8. SKU销售数据分析
9. 搜索词与转化分析
10. 流量来源分析
11. 支付人群画像分析
12. 综合竞品评估
"""

import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _safe_str(value, default="未获取"):
    """安全转换为字符串"""
    if value is None:
        return default
    if isinstance(value, list):
        return "、".join(str(v) for v in value if v) or default
    if isinstance(value, dict):
        return str(value)
    return str(value) if value else default


def _add_section_title(doc, title, level=2):
    """添加章节标题"""
    heading = doc.add_heading(title, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)


def _add_sub_title(doc, title):
    """添加子标题"""
    heading = doc.add_heading(title, level=3)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_key_value_table(doc, data_pairs: list):
    """添加键值对表格"""
    table = doc.add_table(rows=len(data_pairs), cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, (label, value) in enumerate(data_pairs):
        cell_label = table.cell(row_idx, 0)
        cell_value = table.cell(row_idx, 1)

        cell_label.text = str(label)
        cell_value.text = _safe_str(value)

        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()


def _add_paragraph(doc, text, bold=False, font_size=10.5):
    """添加段落文本"""
    if not text:
        return
    para = doc.add_paragraph()
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)


def _add_bullet_list(doc, items, prefix="• "):
    """添加无序列表"""
    if not items:
        return
    if isinstance(items, str):
        items = [items]
    for item in items:
        if item:
            doc.add_paragraph(f"{prefix}{item}", style='List Bullet')


def _section_has_data(data):
    """检查数据是否有效"""
    if not data:
        return False
    if isinstance(data, dict):
        return not data.get("error") and not data.get("_parse_failed")
    return True


def create_analysis_report(analysis_result: dict, output_folder: str) -> str:
    """
    生成竞品分析 Word 报告（9大维度 + 综合评估）

    Args:
        analysis_result: 完整的竞品分析结果
        output_folder: 输出目录（竞品自身的文件夹内）

    Returns:
        str: 生成的文档路径
    """
    doc = Document()

    # 提取各维度数据
    main_strategy = analysis_result.get("main_images_strategy", {})
    detail_logic = analysis_result.get("detail_page_logic", {})
    data_analysis = analysis_result.get("data_analysis", {})
    pdf_analysis = analysis_result.get("pdf_analysis", [])
    competitive_summary = analysis_result.get("competitive_summary", {})

    # 竞品名称
    product_name = (
        main_strategy.get("product_name")
        or analysis_result.get("competitor_name", "未知竞品")
    )

    # ========================================================================
    # 封面
    # ========================================================================
    title = doc.add_heading(f"{product_name}\n竞品深度分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"分析日期：{datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("基于 AI 视觉识别 + 运营数据的多维度竞品分析")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sub_run.italic = True

    doc.add_page_break()

    # ========================================================================
    # 一、产品概况
    # ========================================================================
    _add_section_title(doc, "一、产品概况")

    basic_pairs = [
        ("产品名称", product_name),
        ("分析来源", "电商主图 + 详情页 + 运营数据"),
        ("分析日期", datetime.now().strftime('%Y年%m月%d日')),
    ]
    _add_key_value_table(doc, basic_pairs)

    # ========================================================================
    # 二、5张主图卖点与视觉策略分析
    # ========================================================================
    if _section_has_data(main_strategy):
        _add_section_title(doc, "二、5张主图卖点与视觉策略分析")

        # 逐张主图分析
        images_analysis = main_strategy.get("images_analysis", [])
        if images_analysis:
            _add_sub_title(doc, "2.1 逐张主图分析")

            for img_data in images_analysis:
                if isinstance(img_data, dict):
                    img_no = img_data.get("image_no", "?")
                    doc.add_heading(f"第{img_no}张主图", level=4)

                    img_pairs = [
                        ("🎯 角色定位", img_data.get("role")),
                        ("💡 核心卖点", img_data.get("selling_point")),
                        ("📝 文案策略", img_data.get("copy_strategy")),
                        ("🎨 视觉设计", img_data.get("visual_design")),
                        ("⚖️ 合规评估", img_data.get("compliance_note")),
                    ]
                    _add_key_value_table(doc, img_pairs)

        # 整体评价
        _add_sub_title(doc, "2.2 主图体系整体评价")

        overall_pairs = [
            ("整体策略逻辑", main_strategy.get("overall_strategy")),
            ("卖点排布评价", main_strategy.get("selling_points_layout")),
            ("竞争力评价", main_strategy.get("competitive_rating")),
            ("最佳实践", main_strategy.get("best_practices")),
        ]
        _add_key_value_table(doc, overall_pairs)

        # 改进建议
        suggestions = main_strategy.get("improvement_suggestions", [])
        if suggestions:
            _add_sub_title(doc, "2.3 改进建议")
            _add_bullet_list(doc, suggestions)

        doc.add_paragraph()

    # ========================================================================
    # 三、详情页表达逻辑分析
    # ========================================================================
    if _section_has_data(detail_logic):
        _add_section_title(doc, "三、详情页表达逻辑分析")

        # 内容逻辑线
        logic_flow = detail_logic.get("content_logic_flow")
        if logic_flow:
            _add_sub_title(doc, "3.1 内容逻辑线")
            _add_paragraph(doc, logic_flow)

        # 说服路径
        persuasion = detail_logic.get("persuasion_path")
        if persuasion:
            _add_sub_title(doc, "3.2 核心说服路径")
            _add_paragraph(doc, persuasion)

        # 信任构建
        trust = detail_logic.get("trust_building", {})
        if trust and isinstance(trust, dict):
            _add_sub_title(doc, "3.3 信任构建手段")
            methods = trust.get("methods", [])
            if methods:
                _add_bullet_list(doc, methods)
            evaluation = trust.get("evaluation")
            if evaluation:
                _add_paragraph(doc, f"评价：{evaluation}")

        # 成分与功效表达
        ingredients = detail_logic.get("ingredients_expression", {})
        if ingredients and isinstance(ingredients, dict):
            _add_sub_title(doc, "3.4 成分与功效表达")
            ingredient_pairs = [
                ("核心成分", ingredients.get("core_ingredients")),
                ("功效宣称", ingredients.get("efficacy_claims")),
                ("表达技巧", ingredients.get("expression_technique")),
            ]
            _add_key_value_table(doc, ingredient_pairs)

        # 视觉风格
        visual = detail_logic.get("visual_style", {})
        if visual and isinstance(visual, dict):
            _add_sub_title(doc, "3.5 视觉表达风格")
            style_pairs = [
                ("风格关键词", visual.get("style_keywords")),
                ("定位匹配度", visual.get("positioning_match")),
            ]
            _add_key_value_table(doc, style_pairs)

        # 差异化
        diff = detail_logic.get("differentiation")
        if diff:
            _add_sub_title(doc, "3.6 差异化表达")
            _add_paragraph(doc, diff)

        # 改进建议
        weaknesses = detail_logic.get("weaknesses")
        suggestions = detail_logic.get("improvement_suggestions", [])
        if weaknesses or suggestions:
            _add_sub_title(doc, "3.7 不足与改进建议")
            if weaknesses:
                _add_paragraph(doc, f"不足之处：{weaknesses}")
            if suggestions:
                _add_bullet_list(doc, suggestions)

        doc.add_paragraph()

    # ========================================================================
    # 四~十：数据维度分析
    # ========================================================================
    if _section_has_data(data_analysis):
        # —— 四、规格设计 ——
        spec = data_analysis.get("spec_design", {})
        if isinstance(spec, dict) and spec.get("data_available", True):
            _add_section_title(doc, "四、规格设计分析")
            spec_pairs = [
                ("SKU规格列表", spec.get("sku_list")),
                ("定价策略逻辑", spec.get("pricing_logic")),
                ("规格设计评价", spec.get("evaluation")),
            ]
            _add_key_value_table(doc, spec_pairs)

        # —— 五、活动与促销策略 ——
        promo = data_analysis.get("promotions", {})
        if isinstance(promo, dict) and promo.get("data_available", True):
            _add_section_title(doc, "五、活动与促销策略")
            activities = promo.get("current_activities", [])
            if activities:
                _add_sub_title(doc, "当前参与活动")
                _add_bullet_list(doc, activities)
            promo_pairs = [
                ("优惠券策略", promo.get("coupon_strategy")),
                ("促销力度评价", promo.get("promo_intensity")),
            ]
            _add_key_value_table(doc, promo_pairs)

        # —— 六、用户评价分析 ——
        reviews = data_analysis.get("reviews_analysis", {})
        if isinstance(reviews, dict) and reviews.get("data_available", True):
            _add_section_title(doc, "六、用户评价分析")
            review_pairs = [
                ("评价总数量级", reviews.get("total_volume")),
                ("好评率", reviews.get("positive_rate")),
            ]
            _add_key_value_table(doc, review_pairs)

            pos_kw = reviews.get("positive_keywords", [])
            if pos_kw:
                _add_sub_title(doc, "好评高频关键词")
                _add_bullet_list(doc, pos_kw)

            neg_kw = reviews.get("negative_keywords", [])
            if neg_kw:
                _add_sub_title(doc, "差评/中评关注点")
                _add_bullet_list(doc, neg_kw)

            insights = reviews.get("consumer_insights")
            if insights:
                _add_sub_title(doc, "消费者洞察")
                _add_paragraph(doc, insights)

        # —— 七、近30日销售数据 ——
        sales = data_analysis.get("sales_30days", {})
        if isinstance(sales, dict) and sales.get("data_available", True):
            _add_section_title(doc, "七、近30日销售数据分析")
            sales_pairs = [
                ("近30天销售额/量", sales.get("total_sales")),
                ("日均销售", sales.get("daily_average")),
                ("趋势分析", sales.get("trend_analysis")),
                ("品类定位", sales.get("category_position")),
            ]
            _add_key_value_table(doc, sales_pairs)

        # —— 八、SKU销售数据 ——
        sku = data_analysis.get("sku_sales", {})
        if isinstance(sku, dict) and sku.get("data_available", True):
            _add_section_title(doc, "八、SKU销售数据分析")
            sku_pairs = [
                ("最畅销SKU", sku.get("top_sku")),
                ("各SKU销售分布", sku.get("sales_distribution")),
                ("策略洞察", sku.get("strategy_insight")),
            ]
            _add_key_value_table(doc, sku_pairs)

        # —— 九、搜索词与转化 ——
        search = data_analysis.get("search_terms", {})
        if isinstance(search, dict) and search.get("data_available", True):
            _add_section_title(doc, "九、搜索词与转化分析")

            top_terms = search.get("top_terms", [])
            if top_terms:
                _add_sub_title(doc, "主要搜索词")
                _add_bullet_list(doc, top_terms)

            search_pairs = [
                ("转化率分析", search.get("conversion_analysis")),
                ("关键词策略", search.get("keyword_strategy")),
                ("品牌词vs品类词", search.get("brand_vs_category")),
            ]
            _add_key_value_table(doc, search_pairs)

        # —— 十、流量来源 ——
        traffic = data_analysis.get("traffic_sources", {})
        if isinstance(traffic, dict) and traffic.get("data_available", True):
            _add_section_title(doc, "十、流量来源分析")
            traffic_pairs = [
                ("渠道流量分布", traffic.get("channel_distribution")),
                ("主要流量来源", traffic.get("primary_source")),
                ("付费/免费比例", traffic.get("paid_vs_organic")),
                ("流量健康度评估", traffic.get("health_assessment")),
            ]
            _add_key_value_table(doc, traffic_pairs)

    # ========================================================================
    # 十一、支付人群画像分析（来自 PDF 视觉识别）
    # ========================================================================
    if pdf_analysis and isinstance(pdf_analysis, list):
        _add_section_title(doc, "十一、支付人群画像分析")

        for pdf_idx, pdf_item in enumerate(pdf_analysis):
            if not isinstance(pdf_item, dict) or pdf_item.get("_parse_failed"):
                continue

            pdf_source = pdf_item.get("pdf_source", "PDF文件")
            content_type = pdf_item.get("content_type", "未知")

            if len(pdf_analysis) > 1:
                _add_sub_title(doc, f"11.{pdf_idx+1} {pdf_source}")

            # 内容类型和摘要
            summary = pdf_item.get("summary", "")
            if content_type or summary:
                basic_pairs = [
                    ("数据来源", pdf_source),
                    ("内容类型", content_type),
                ]
                if summary:
                    basic_pairs.append(("核心摘要", summary))
                _add_key_value_table(doc, basic_pairs)

            # 关键数据
            key_data = pdf_item.get("key_data", {})
            if key_data and isinstance(key_data, dict):
                _add_sub_title(doc, "关键数据指标")
                kd_pairs = []
                for k, v in key_data.items():
                    kd_pairs.append((k, v))
                _add_key_value_table(doc, kd_pairs)

            # 图表数据
            charts = pdf_item.get("charts_data", [])
            if charts and isinstance(charts, list):
                _add_sub_title(doc, "图表数据详情")
                for chart in charts:
                    if isinstance(chart, dict):
                        chart_title = chart.get("chart_title", "图表")
                        chart_type = chart.get("chart_type", "")
                        data_points = chart.get("data_points", "")

                        doc.add_heading(f"{chart_title}（{chart_type}）", level=4)

                        # 处理不同格式的 data_points
                        if isinstance(data_points, dict):
                            dp_pairs = [(k, v) for k, v in data_points.items()]
                            _add_key_value_table(doc, dp_pairs)
                        elif isinstance(data_points, list):
                            for dp in data_points:
                                if isinstance(dp, dict):
                                    dp_text = " | ".join(f"{k}: {v}" for k, v in dp.items())
                                    doc.add_paragraph(f"• {dp_text}", style='List Bullet')
                                else:
                                    doc.add_paragraph(f"• {dp}", style='List Bullet')
                        elif data_points:
                            _add_paragraph(doc, str(data_points))

            # 专业解读
            insights = pdf_item.get("professional_insights", "")
            if insights:
                _add_sub_title(doc, "专业解读")
                _add_paragraph(doc, insights)

        doc.add_paragraph()

    # ========================================================================
    # 十二、综合竞品评估
    # ========================================================================
    if _section_has_data(competitive_summary):
        doc.add_page_break()
        _add_section_title(doc, "十二、综合竞品评估与策略建议")

        # 优势
        strengths = competitive_summary.get("competitor_strengths", [])
        if strengths:
            _add_sub_title(doc, "🔴 核心竞争优势")
            if isinstance(strengths, list):
                _add_bullet_list(doc, strengths)
            else:
                _add_paragraph(doc, str(strengths))

        # 劣势
        weaknesses = competitive_summary.get("competitor_weaknesses", [])
        if weaknesses:
            _add_sub_title(doc, "🟢 核心竞争劣势（我方突破口）")
            if isinstance(weaknesses, list):
                _add_bullet_list(doc, weaknesses)
            else:
                _add_paragraph(doc, str(weaknesses))

        # 差异化机会
        opportunities = competitive_summary.get("differentiation_opportunities", [])
        if opportunities:
            _add_sub_title(doc, "💡 差异化机会点")
            if isinstance(opportunities, list):
                for opp in opportunities:
                    if isinstance(opp, dict):
                        direction = opp.get("direction", "")
                        how_to = opp.get("how_to", "")
                        doc.add_paragraph(f"• {direction}", style='List Bullet')
                        if how_to:
                            _add_paragraph(doc, f"  操作建议：{how_to}")
                    else:
                        doc.add_paragraph(f"• {opp}", style='List Bullet')
            else:
                _add_paragraph(doc, str(opportunities))

        # 竞争策略
        strategy = competitive_summary.get("competition_strategy", {})
        if strategy and isinstance(strategy, dict):
            _add_sub_title(doc, "🎯 竞争策略建议")
            strategy_pairs = [
                ("产品力策略", strategy.get("product")),
                ("营销力策略", strategy.get("marketing")),
                ("渠道力策略", strategy.get("channel")),
            ]
            _add_key_value_table(doc, strategy_pairs)

        # 威胁评级
        threat = competitive_summary.get("overall_threat_level")
        if threat:
            _add_sub_title(doc, "⚡ 竞争威胁评级")
            _add_paragraph(doc, threat)

        # 深入分析建议
        further = competitive_summary.get("further_analysis_suggestions", [])
        if further:
            _add_sub_title(doc, "📌 建议深入分析方向")
            _add_bullet_list(doc, further)
            _add_paragraph(doc, "\n（可针对以上方向继续对话，获取更深入的分析）", font_size=9)

    # ========================================================================
    # 保存文档
    # ========================================================================
    output_dir = Path(output_folder)
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_name = product_name.replace("/", "_").replace("\\", "_").replace(":", "_")
    doc_filename = f"{safe_name}_竞品分析.docx"
    doc_path = output_dir / doc_filename

    doc.save(str(doc_path))
    print(f"  ✓ Word 报告: {doc_path}")

    return str(doc_path)


def create_batch_reports(all_results: list, competitors_base_path: str) -> list:
    """
    批量生成多个竞品的分析报告
    每个竞品的报告保存在自己的文件夹内

    Args:
        all_results: 所有竞品的分析结果列表
        competitors_base_path: 竞品素材基础目录

    Returns:
        list[str]: 生成的文档路径列表
    """
    doc_paths = []
    for result in all_results:
        if "error" in result and not result.get("main_images_strategy"):
            print(f"  ⚠ 跳过 {result.get('competitor_name', '未知')}: 分析失败")
            continue
        try:
            # 关键：保存到竞品自身的文件夹内，而非基础目录
            comp_name = result.get("competitor_name", "未知竞品")
            comp_folder = Path(competitors_base_path) / comp_name
            if not comp_folder.exists():
                # 如果竞品文件夹不存在，才回退到基础目录
                comp_folder = Path(competitors_base_path)

            path = create_analysis_report(result, str(comp_folder))
            doc_paths.append(path)
        except Exception as e:
            print(f"  ✗ 生成报告失败: {e}")

    return doc_paths
